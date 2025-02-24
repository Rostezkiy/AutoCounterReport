import customtkinter as ctk
from tkinter import messagebox
import configparser
from docx import Document
import locale
from datetime import datetime
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders


class WaterReportApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Программа учета воды")
        self.geometry("800x600")
        self.resizable(False, False)
        self.tab_view = ctk.CTkTabview(self)
        self.tab_view.pack(fill="both", expand=True, padx=10, pady=(10, 0))
        self.tab_view.add("Ввод данных")
        self.tab_view.add("Настройки SMTP")
        self.create_input_tab()
        self.create_settings_tab()

    def create_input_tab(self):
        input_frame = ctk.CTkFrame(self.tab_view.tab("Ввод данных"))
        input_frame.pack(fill="both", expand=True, padx=20, pady=20)
        self.months_ru = {
            'January': 'января',
            'February': 'февраля',
            'March': 'марта',
            'April': 'апреля',
            'May': 'мая',
            'June': 'июня',
            'July': 'июля',
            'August': 'августа',
            'September': 'сентября',
            'October': 'октября',
            'November': 'ноября',
            'December': 'декабря'
        }
        current_date = datetime.now()

        date_text = f'«{current_date.day}» {self.months_ru[current_date.strftime("%B")]} {current_date.year}г'

        # Поле даты
        ctk.CTkLabel(input_frame, text="Текущая дата:").pack(anchor="w")
        self.date_label = ctk.CTkLabel(input_frame,
                                       text=date_text)
        self.date_label.pack(anchor="w")

        # Поле холодной воды
        ctk.CTkLabel(input_frame, text="Холодная вода (куб.м):").pack(anchor="w", pady=(20, 5))
        self.cold_water_entry = ctk.CTkEntry(input_frame)
        self.cold_water_entry.pack(fill="x")

        # Поле горячей воды
        ctk.CTkLabel(input_frame, text="Горячая вода (куб.м):").pack(anchor="w", pady=(20, 5))
        self.hot_water_entry = ctk.CTkEntry(input_frame)
        self.hot_water_entry.pack(fill="x")

        # Кнопка формирования отчета
        self.generate_button = ctk.CTkButton(
            input_frame,
            text="Сформировать отчет",
            command=self.generate_report
        )
        self.generate_button.pack(fill="x", pady=(20, 0))

    def create_settings_tab(self):
        # Вкладка настроек
        settings_frame = ctk.CTkFrame(self.tab_view.tab("Настройки SMTP"))
        settings_frame.pack(fill="both", expand=True, padx=20, pady=20)

        # Поле сервера
        ctk.CTkLabel(settings_frame, text="Сервер SMTP:").pack(anchor="w")
        self.smtp_server_entry = ctk.CTkEntry(settings_frame)
        self.smtp_server_entry.pack(fill="x")

        # Поле порта
        ctk.CTkLabel(settings_frame, text="Порт:").pack(anchor="w", pady=(10, 5))
        self.port_entry = ctk.CTkEntry(settings_frame)
        self.port_entry.pack(fill="x")

        # Поле email от
        ctk.CTkLabel(settings_frame, text="Email от:").pack(anchor="w", pady=(10, 5))
        self.email_from_entry = ctk.CTkEntry(settings_frame)
        self.email_from_entry.pack(fill="x")

        # Поле email кому
        ctk.CTkLabel(settings_frame, text="Email кому:").pack(anchor="w", pady=(10, 5))
        self.email_to_entry = ctk.CTkEntry(settings_frame)
        self.email_to_entry.pack(fill="x")

        # Поле пароля
        ctk.CTkLabel(settings_frame, text="Пароль:").pack(anchor="w", pady=(10, 5))
        self.password_entry = ctk.CTkEntry(settings_frame, show="*")
        self.password_entry.pack(fill="x")

        # Кнопка сохранения настроек
        self.save_settings_button = ctk.CTkButton(
            settings_frame,
            text="Сохранить настройки",
            command=self.save_settings
        )
        self.save_settings_button.pack(fill="x", pady=(20, 0))

        # Загрузка существующих настроек
        self.load_settings()

    def load_settings(self):
        config = configparser.ConfigParser()
        try:
            config.read('config.ini')
            self.smtp_server_entry.insert(0, config.get('SMTP', 'server'))
            self.port_entry.insert(0, config.get('SMTP', 'port'))
            self.email_from_entry.insert(0, config.get('SMTP', 'email_from'))
            self.email_to_entry.insert(0, config.get('SMTP', 'email_to'))
            self.password_entry.insert(0, config.get('SMTP', 'password'))

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при загрузке настроек: {e}")

    def save_settings(self):
        config = configparser.ConfigParser()
        config['SMTP'] = {
            'server': self.smtp_server_entry.get(),
            'port': self.port_entry.get(),
            'email_from': self.email_from_entry.get(),
            'email_to': self.email_to_entry.get(),
            'password': self.password_entry.get()
        }

        try:
            with open('config.ini', 'w') as configfile:
                config.write(configfile)
            messagebox.showinfo("Успех", "Настройки успешно сохранены!")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить настройки: {str(e)}")

    def generate_report(self):
        try:
            locale.setlocale(locale.LC_ALL, 'ru_RU.UTF-8')
        except locale.Error:
            messagebox.showerror("Ошибка локали", "Локаль 'ru_RU.UTF-8' не поддерживается.")
            return

        date = datetime.now().strftime('«%d» %B %Yг')
        filename = f"Отчет_{datetime.now().strftime('%d %B %Yг')}.docx".replace(" ", "_")

        try:
            doc = Document('Отчет.docx')
            paragraphs = doc.paragraphs
            table = doc.tables[0]

            cold_water_value = float(self.cold_water_entry.get())
            hot_water_value = float(self.hot_water_entry.get())

            for i in range(len(table.columns)):
                table.cell(1, i).text = table.cell(2, i).text

            table.cell(2, 0).text = datetime.now().strftime('%d %B')
            table.cell(2, 1).text = f"{cold_water_value} куб. м."
            table.cell(2, 2).text = f"{hot_water_value} куб. м."

            cold_water_1 = float(table.cell(1, 1).text.split()[0])
            hot_water_1 = float(table.cell(1, 2).text.split()[0])

            table.cell(3, 1).text = f'{round(cold_water_value - cold_water_1, 2)} куб. м.'
            table.cell(3, 2).text = f'{round(hot_water_value - hot_water_1, 2)} куб. м.'

            paragraphs[-2].text = f'Подпись квартиросъемщика_____________________________________Дата {date}'
            paragraphs[-1].text = f'Сведения принял______________________________________________Дата {date}'
            doc.save(filename)

            if messagebox.askyesno(
                    "Отчет сформирован",
                    "Отчет успешно сформирован! Отправить его на почту?"
            ):
                self.send_email(filename)

        except FileNotFoundError:
            messagebox.showerror("Ошибка", "Файл 'Отчет.docx' не найден!")
        except ValueError:
            messagebox.showerror("Ошибка", "Введите корректные числовые значения!")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Произошла ошибка: {str(e)}")

    def send_email(self, filename):
        config = configparser.ConfigParser()
        try:
            config.read('config.ini')
            server = config.get('SMTP', 'server')
            port = int(config.get('SMTP', 'port'))
            email_from = config.get('SMTP', 'email_from')
            email_to = config.get('SMTP', 'email_to')
            password = config.get('SMTP', 'password')

            msg = MIMEMultipart()
            msg['From'] = email_from
            msg['To'] = email_to
            msg['Subject'] = 'Передача данных'

            with open(filename, "rb") as attachment:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition',
                                'attachment',
                                filename=('utf-8', '', filename))
                msg.attach(part)

            with smtplib.SMTP(server, port) as smtp:
                smtp.starttls()
                smtp.login(email_from, password)
                smtp.send_message(msg)

            messagebox.showinfo("Успех", "Письмо успешно отправлено!")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось отправить письмо: {str(e)}")


if __name__ == "__main__":
    app = WaterReportApp()
    app.mainloop()